"""Export meeting minutes to DOCX and PDF."""

import logging
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak

from src.config import Config
from src.schemas import Meeting, ExportOptions

logger = logging.getLogger(__name__)


class DocumentExporter:
    """Export meeting minutes to various formats."""

    def __init__(self):
        """Initialize exporter."""
        self.exports_dir = Config.EXPORTS_DIR
        self.exports_dir.mkdir(exist_ok=True)

    def export(
        self,
        meeting: Meeting,
        options: ExportOptions
    ) -> Path:
        """Export meeting to specified format.

        Args:
            meeting: Meeting object
            options: Export options

        Returns:
            Path to exported file
        """
        if options.format == "docx":
            return self.export_docx(meeting, options)
        elif options.format == "pdf":
            return self.export_pdf(meeting, options)
        else:
            raise ValueError(f"Unsupported format: {options.format}")

    def export_docx(
        self,
        meeting: Meeting,
        options: ExportOptions
    ) -> Path:
        """Export to DOCX format.

        Args:
            meeting: Meeting object
            options: Export options

        Returns:
            Path to DOCX file
        """
        logger.info(f"Exporting meeting to DOCX: {meeting.title}")

        # Create document
        doc = Document()

        # Add title
        title = doc.add_heading("Meeting Minutes", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata
        doc.add_heading("Meeting Information", 1)

        metadata_table = doc.add_table(rows=4, cols=2)
        metadata_table.style = 'Light Grid Accent 1'

        # Title
        metadata_table.rows[0].cells[0].text = "Title:"
        metadata_table.rows[0].cells[1].text = meeting.title

        # Date
        metadata_table.rows[1].cells[0].text = "Date:"
        metadata_table.rows[1].cells[1].text = meeting.date.strftime("%B %d, %Y %H:%M")

        # Participants
        metadata_table.rows[2].cells[0].text = "Participants:"
        metadata_table.rows[2].cells[1].text = ", ".join(meeting.participants) if meeting.participants else "N/A"

        # Agenda
        metadata_table.rows[3].cells[0].text = "Agenda:"
        metadata_table.rows[3].cells[1].text = meeting.agenda or "N/A"

        doc.add_paragraph()

        if meeting.minutes:
            # Executive Summary
            if meeting.minutes.summary:
                doc.add_heading("Executive Summary", 1)
                for bullet in meeting.minutes.summary:
                    doc.add_paragraph(bullet, style='List Bullet')
                doc.add_paragraph()

            # Key Decisions
            if meeting.minutes.decisions:
                doc.add_heading("Key Decisions", 1)
                for decision in meeting.minutes.decisions:
                    doc.add_paragraph(decision, style='List Bullet')
                doc.add_paragraph()

            # Action Items
            if meeting.minutes.action_items:
                doc.add_heading("Action Items", 1)

                # Create action items table
                action_table = doc.add_table(rows=len(meeting.minutes.action_items) + 1, cols=4)
                action_table.style = 'Light Grid Accent 1'

                # Header row
                headers = ["Owner", "Action", "Due Date", "Status"]
                for i, header in enumerate(headers):
                    cell = action_table.rows[0].cells[i]
                    cell.text = header
                    # Make header bold
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

                # Data rows
                for i, item in enumerate(meeting.minutes.action_items, 1):
                    action_table.rows[i].cells[0].text = item.owner
                    action_table.rows[i].cells[1].text = item.task
                    action_table.rows[i].cells[2].text = item.due_date or "TBD"
                    action_table.rows[i].cells[3].text = item.status

                doc.add_paragraph()

            # Risks and Open Questions
            if meeting.minutes.risks:
                doc.add_heading("Risks & Open Questions", 1)
                for risk in meeting.minutes.risks:
                    doc.add_paragraph(risk, style='List Bullet')
                doc.add_paragraph()

            # Additional Notes
            if meeting.minutes.notes:
                doc.add_heading("Additional Notes", 1)
                doc.add_paragraph(meeting.minutes.notes)
                doc.add_paragraph()

        # Transcript (if included)
        if options.include_transcript and meeting.transcript:
            doc.add_page_break()
            doc.add_heading("Meeting Transcript", 1)

            current_speaker = None
            for seg in meeting.transcript:
                if options.include_speaker_labels and seg.speaker != current_speaker:
                    # Add speaker header
                    speaker_para = doc.add_paragraph()
                    speaker_run = speaker_para.add_run(f"\n{seg.speaker}")
                    speaker_run.bold = True
                    speaker_run.font.color.rgb = RGBColor(0, 102, 204)
                    current_speaker = seg.speaker

                # Add transcript text
                if options.include_timestamps:
                    timestamp = self._format_timestamp(seg.start)
                    text = f"[{timestamp}] {seg.text}"
                else:
                    text = seg.text

                doc.add_paragraph(text)

        # Footer
        doc.add_page_break()
        footer_para = doc.add_paragraph()
        footer_para.add_run(f"Generated by Meeting Minutes Generator on {datetime.now().strftime('%B %d, %Y %H:%M')}")
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.runs[0].font.size = Pt(9)
        footer_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

        # Save document
        filename = self._generate_filename(meeting, "docx")
        filepath = self.exports_dir / filename
        doc.save(filepath)

        logger.info(f"DOCX exported: {filepath}")
        return filepath

    def export_pdf(
        self,
        meeting: Meeting,
        options: ExportOptions
    ) -> Path:
        """Export to PDF format.

        Args:
            meeting: Meeting object
            options: Export options

        Returns:
            Path to PDF file
        """
        logger.info(f"Exporting meeting to PDF: {meeting.title}")

        filename = self._generate_filename(meeting, "pdf")
        filepath = self.exports_dir / filename

        # Create PDF document
        doc = SimpleDocTemplate(
            str(filepath),
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )

        # Container for elements
        story = []

        # Styles
        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=24,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=30,
            alignment=1  # Center
        )

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=colors.HexColor('#0066cc'),
            spaceAfter=12,
            spaceBefore=12
        )

        # Title
        story.append(Paragraph("Meeting Minutes", title_style))
        story.append(Spacer(1, 0.2*inch))

        # Metadata table
        metadata_data = [
            ["Title:", meeting.title],
            ["Date:", meeting.date.strftime("%B %d, %Y %H:%M")],
            ["Participants:", ", ".join(meeting.participants) if meeting.participants else "N/A"],
            ["Agenda:", meeting.agenda or "N/A"]
        ]

        metadata_table = Table(metadata_data, colWidths=[1.5*inch, 5*inch])
        metadata_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f0f0')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey)
        ]))

        story.append(metadata_table)
        story.append(Spacer(1, 0.3*inch))

        if meeting.minutes:
            # Executive Summary
            if meeting.minutes.summary:
                story.append(Paragraph("Executive Summary", heading_style))
                for bullet in meeting.minutes.summary:
                    story.append(Paragraph(f"• {bullet}", styles['Normal']))
                    story.append(Spacer(1, 0.1*inch))
                story.append(Spacer(1, 0.2*inch))

            # Key Decisions
            if meeting.minutes.decisions:
                story.append(Paragraph("Key Decisions", heading_style))
                for decision in meeting.minutes.decisions:
                    story.append(Paragraph(f"• {decision}", styles['Normal']))
                    story.append(Spacer(1, 0.1*inch))
                story.append(Spacer(1, 0.2*inch))

            # Action Items
            if meeting.minutes.action_items:
                story.append(Paragraph("Action Items", heading_style))

                action_data = [["Owner", "Action", "Due Date", "Status"]]
                for item in meeting.minutes.action_items:
                    action_data.append([
                        item.owner,
                        item.task,
                        item.due_date or "TBD",
                        item.status
                    ])

                action_table = Table(action_data, colWidths=[1.2*inch, 3*inch, 1*inch, 0.8*inch])
                action_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0066cc')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 11),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))

                story.append(action_table)
                story.append(Spacer(1, 0.2*inch))

            # Risks
            if meeting.minutes.risks:
                story.append(Paragraph("Risks & Open Questions", heading_style))
                for risk in meeting.minutes.risks:
                    story.append(Paragraph(f"• {risk}", styles['Normal']))
                    story.append(Spacer(1, 0.1*inch))
                story.append(Spacer(1, 0.2*inch))

        # Transcript (if included)
        if options.include_transcript and meeting.transcript:
            story.append(PageBreak())
            story.append(Paragraph("Meeting Transcript", heading_style))

            current_speaker = None
            for seg in meeting.transcript:
                if options.include_speaker_labels and seg.speaker != current_speaker:
                    speaker_style = ParagraphStyle(
                        'Speaker',
                        parent=styles['Normal'],
                        fontName='Helvetica-Bold',
                        textColor=colors.HexColor('#0066cc'),
                        spaceAfter=6,
                        spaceBefore=12
                    )
                    story.append(Paragraph(seg.speaker, speaker_style))
                    current_speaker = seg.speaker

                if options.include_timestamps:
                    timestamp = self._format_timestamp(seg.start)
                    text = f"[{timestamp}] {seg.text}"
                else:
                    text = seg.text

                story.append(Paragraph(text, styles['Normal']))
                story.append(Spacer(1, 0.05*inch))

        # Footer
        footer_text = f"Generated by Meeting Minutes Generator on {datetime.now().strftime('%B %d, %Y %H:%M')}"
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=8,
            textColor=colors.grey,
            alignment=1
        )
        story.append(Spacer(1, 0.5*inch))
        story.append(Paragraph(footer_text, footer_style))

        # Build PDF
        doc.build(story)

        logger.info(f"PDF exported: {filepath}")
        return filepath

    def _generate_filename(self, meeting: Meeting, extension: str) -> str:
        """Generate filename for export.

        Args:
            meeting: Meeting object
            extension: File extension (docx or pdf)

        Returns:
            Filename string
        """
        # Sanitize title for filename
        safe_title = "".join(c for c in meeting.title if c.isalnum() or c in (' ', '-', '_'))
        safe_title = safe_title.replace(' ', '_')[:50]

        date_str = meeting.date.strftime("%Y%m%d")
        timestamp = datetime.now().strftime("%H%M%S")

        return f"{date_str}_{safe_title}_{timestamp}.{extension}"

    def _format_timestamp(self, seconds: float) -> str:
        """Format timestamp as MM:SS.

        Args:
            seconds: Time in seconds

        Returns:
            Formatted timestamp
        """
        minutes = int(seconds // 60)
        secs = int(seconds % 60)
        return f"{minutes:02d}:{secs:02d}"
